#!/usr/bin/env python3
"""
Report Templates

Enhanced reporting with formal document templates for:
- PDF export via ReportLab
- Word export via python-docx
- Governance gate reports
- Executive summaries
- Custom report builder

Author: Portfolio ML
Version: 1.0.0
"""

from typing import Dict, List, Optional
from datetime import datetime
import os


class ReportExporter:
    """
    Export project plans and team recommendations to multiple formats
    
    Supports: PDF, Word, HTML
    """
    
    def __init__(self):
        """Initialize report exporter"""
        self.templates = {
            'standard': 'Standard Project Plan',
            'executive': 'Executive Summary',
            'governance_gate': 'Governance Gate Report',
            'team_composition': 'Team Composition Report'
        }
    
    def export_to_pdf(self, plan, output_path: str, template: str = 'standard') -> str:
        """
        Export project plan to PDF
        
        Args:
            plan: ProjectPlan object
            output_path: Output file path
            template: Report template to use
        
        Returns:
            Output file path
        
        Note: Requires reportlab: pip install reportlab
        """
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                PageBreak, Image
            )
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        except ImportError:
            # Fallback to markdown if reportlab not available
            print("⚠️  reportlab not installed. Install with: pip install reportlab")
            print("   Falling back to markdown export...")
            from project_plan_generator import ProjectPlanGenerator
            generator = ProjectPlanGenerator()
            md_path = output_path.replace('.pdf', '.md')
            return generator.export_to_markdown(plan, md_path)
        
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            leftMargin=1*inch,
            rightMargin=1*inch
        )
        
        # Container for content
        content = []
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=12,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=12,
            spaceBefore=12
        )
        normal_style = styles['Normal']
        
        # Title page
        content.append(Paragraph(f"Project Plan: {plan.charter.project_name}", title_style))
        content.append(Spacer(1, 0.2*inch))
        
        meta_data = [
            ['Project ID:', plan.charter.project_id],
            ['Generated:', plan.generated_date.strftime('%Y-%m-%d %H:%M')],
            ['Duration:', f"{plan.timeline['duration_months']} months"],
            ['Budget:', f"${plan.budget['total_cost']:,.0f}"],
            ['ROI:', f"{plan.budget['financial_summary']['roi_percent']:.1f}%"]
        ]
        
        meta_table = Table(meta_data, colWidths=[2*inch, 4*inch])
        meta_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        content.append(meta_table)
        content.append(PageBreak())
        
        # Executive Summary
        content.append(Paragraph("Executive Summary", heading_style))
        content.append(Paragraph(plan.charter.executive_summary, normal_style))
        content.append(Spacer(1, 0.2*inch))
        
        # Objectives
        content.append(Paragraph("Objectives", heading_style))
        for obj in plan.charter.objectives:
            content.append(Paragraph(f"• {obj}", normal_style))
        content.append(Spacer(1, 0.2*inch))
        
        # Strategic Alignment
        content.append(Paragraph("Strategic Alignment", heading_style))
        sa = plan.charter.strategic_alignment
        content.append(Paragraph(
            f"Overall Score: {sa['alignment_score']:.1f}/100 ({sa['alignment_level']})",
            normal_style
        ))
        content.append(Spacer(1, 0.1*inch))
        
        # Timeline
        content.append(PageBreak())
        content.append(Paragraph("Project Timeline", heading_style))
        
        phase_data = [['Phase', 'Duration', 'Start', 'End']]
        for phase in plan.timeline['phases']:
            phase_data.append([
                phase['name'],
                f"{phase['duration_months']} months",
                f"Month {phase['start_month']}",
                f"Month {phase['end_month']}"
            ])
        
        phase_table = Table(phase_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 1*inch])
        phase_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1f4788')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        content.append(phase_table)
        content.append(Spacer(1, 0.2*inch))
        
        # Risks
        content.append(Paragraph("Top Risks", heading_style))
        risk_data = [['ID', 'Category', 'Score', 'Probability', 'Impact']]
        for risk in plan.risk_register[:5]:
            risk_data.append([
                risk['risk_id'],
                risk['category'],
                str(risk['risk_score']),
                risk['probability'],
                risk['impact']
            ])
        
        risk_table = Table(risk_data, colWidths=[0.8*inch, 1.5*inch, 0.8*inch, 1*inch, 1*inch])
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#c0392b')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        content.append(risk_table)
        
        # Build PDF
        doc.build(content)
        
        return output_path
    
    def export_to_word(self, plan, output_path: str, template: str = 'standard') -> str:
        """
        Export project plan to Word
        
        Args:
            plan: ProjectPlan object
            output_path: Output file path
            template: Report template to use
        
        Returns:
            Output file path
        
        Note: Requires python-docx: pip install python-docx
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # Fallback to markdown if python-docx not available
            print("⚠️  python-docx not installed. Install with: pip install python-docx")
            print("   Falling back to markdown export...")
            from project_plan_generator import ProjectPlanGenerator
            generator = ProjectPlanGenerator()
            md_path = output_path.replace('.docx', '.md')
            return generator.export_to_markdown(plan, md_path)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading(f"Project Plan: {plan.charter.project_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        cells = [
            ('Project ID:', plan.charter.project_id),
            ('Generated:', plan.generated_date.strftime('%Y-%m-%d %H:%M')),
            ('Duration:', f"{plan.timeline['duration_months']} months"),
            ('Budget:', f"${plan.budget['total_cost']:,.0f}"),
            ('ROI:', f"{plan.budget['financial_summary']['roi_percent']:.1f}%")
        ]
        
        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            # Bold labels
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(plan.charter.executive_summary)
        
        # Objectives
        doc.add_heading('Objectives', 1)
        for obj in plan.charter.objectives:
            doc.add_paragraph(obj, style='List Bullet')
        
        # Strategic Alignment
        doc.add_heading('Strategic Alignment', 1)
        sa = plan.charter.strategic_alignment
        doc.add_paragraph(
            f"Overall Score: {sa['alignment_score']:.1f}/100 ({sa['alignment_level']})"
        )
        
        doc.add_paragraph('Pillar Scores:')
        for pillar, score in sa['pillar_scores'].items():
            pillar_name = pillar.replace('_', ' ').title()
            doc.add_paragraph(f"{pillar_name}: {score:.1f}/100", style='List Bullet')
        
        # Timeline
        doc.add_page_break()
        doc.add_heading('Project Timeline', 1)
        
        phase_table = doc.add_table(rows=len(plan.timeline['phases'])+1, cols=4)
        phase_table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = phase_table.rows[0].cells
        header_cells[0].text = 'Phase'
        header_cells[1].text = 'Duration'
        header_cells[2].text = 'Start'
        header_cells[3].text = 'End'
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Data
        for i, phase in enumerate(plan.timeline['phases'], 1):
            cells = phase_table.rows[i].cells
            cells[0].text = phase['name']
            cells[1].text = f"{phase['duration_months']} months"
            cells[2].text = f"Month {phase['start_month']}"
            cells[3].text = f"Month {phase['end_month']}"
        
        # Milestones
        doc.add_heading('Milestones & Governance Gates', 1)
        for milestone in plan.milestones:
            gate_marker = " 🚪 GOVERNANCE GATE" if milestone.governance_gate else ""
            doc.add_heading(
                f"Month {milestone.target_date_month}: {milestone.name}{gate_marker}",
                2
            )
            doc.add_paragraph(milestone.description)
            
            if milestone.governance_gate and milestone.gate_criteria:
                doc.add_paragraph('Gate Criteria:')
                for criteria in milestone.gate_criteria:
                    doc.add_paragraph(criteria, style='List Bullet')
        
        # Risks
        doc.add_page_break()
        doc.add_heading('Risk Register', 1)
        
        for risk in plan.risk_register:
            doc.add_heading(f"{risk['risk_id']}: {risk['category']} Risk", 2)
            doc.add_paragraph(f"Description: {risk['description']}")
            doc.add_paragraph(
                f"Probability: {risk['probability']} | Impact: {risk['impact']} | Score: {risk['risk_score']}"
            )
            doc.add_paragraph(f"Mitigation: {risk['mitigation']}")
        
        # Budget
        doc.add_page_break()
        doc.add_heading('Budget & Financial Analysis', 1)
        doc.add_paragraph(f"Total Cost: ${plan.budget['total_cost']:,.0f}")
        
        doc.add_paragraph('Cost Breakdown:')
        for category, amount in plan.budget['cost_breakdown'].items():
            doc.add_paragraph(f"{category}: ${amount:,.0f}", style='List Bullet')
        
        doc.add_paragraph('\nFinancial Metrics:')
        fs = plan.budget['financial_summary']
        metrics = [
            f"NPV: ${fs['npv']:,.0f}",
            f"ROI: {fs['roi_percent']:.1f}%",
            f"Payback Period: {fs['payback_years']:.1f} years",
            f"Benefit/Cost Ratio: {fs['benefit_cost_ratio']:.2f}"
        ]
        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')
        
        # Stakeholders
        doc.add_heading('Stakeholders', 1)
        for stakeholder in plan.stakeholders:
            doc.add_heading(f"{stakeholder.name} ({stakeholder.role})", 2)
            doc.add_paragraph(f"Responsibility: {stakeholder.responsibility}")
            doc.add_paragraph(f"Engagement Level: {stakeholder.engagement_level}")
        
        # Save document
        doc.save(output_path)
        
        return output_path
    
    def export_governance_gate_report(
        self,
        plan,
        milestone_index: int,
        output_path: str,
        format: str = 'pdf'
    ) -> str:
        """
        Export governance gate report
        
        Args:
            plan: ProjectPlan object
            milestone_index: Index of milestone to report on
            output_path: Output file path
            format: 'pdf' or 'word'
        
        Returns:
            Output file path
        """
        milestone = plan.milestones[milestone_index]
        
        if not milestone.governance_gate:
            raise ValueError(f"Milestone {milestone_index} is not a governance gate")
        
        # Create simplified plan focused on gate
        gate_data = {
            'project_name': plan.charter.project_name,
            'project_id': plan.charter.project_id,
            'gate_name': milestone.name,
            'gate_month': milestone.target_date_month,
            'gate_criteria': milestone.gate_criteria,
            'deliverables': milestone.deliverables,
            'current_status': 'PENDING REVIEW',
            'risks': plan.risk_register[:3],  # Top 3 risks
            'budget_status': plan.budget,
            'generated_date': datetime.now()
        }
        
        # For now, export as markdown with gate focus
        md_content = self._generate_gate_markdown(gate_data)
        
        md_path = output_path.replace(f'.{format}', '.md')
        with open(md_path, 'w') as f:
            f.write(md_content)
        
        return md_path
    
    def _generate_gate_markdown(self, gate_data: Dict) -> str:
        """Generate markdown for governance gate report"""
        
        md = f"# Governance Gate Report\n\n"
        md += f"**Project:** {gate_data['project_name']} ({gate_data['project_id']})\n\n"
        md += f"**Gate:** {gate_data['gate_name']}\n\n"
        md += f"**Target Month:** {gate_data['gate_month']}\n\n"
        md += f"**Generated:** {gate_data['generated_date'].strftime('%Y-%m-%d %H:%M')}\n\n"
        md += "---\n\n"
        
        md += "## Gate Criteria\n\n"
        md += "**Status:** " + gate_data['current_status'] + "\n\n"
        
        for i, criteria in enumerate(gate_data['gate_criteria'], 1):
            md += f"{i}. ☐ {criteria}\n"
        md += "\n"
        
        md += "## Expected Deliverables\n\n"
        for deliverable in gate_data['deliverables']:
            md += f"- {deliverable}\n"
        md += "\n"
        
        md += "## Current Risks\n\n"
        for risk in gate_data['risks']:
            md += f"### {risk['risk_id']}: {risk['category']}\n"
            md += f"**Score:** {risk['risk_score']}/100\n\n"
            md += f"{risk['description']}\n\n"
        
        md += "## Financial Status\n\n"
        fs = gate_data['budget_status']['financial_summary']
        md += f"- **Total Cost:** ${gate_data['budget_status']['total_cost']:,.0f}\n"
        md += f"- **NPV:** ${fs['npv']:,.0f}\n"
        md += f"- **ROI:** {fs['roi_percent']:.1f}%\n\n"
        
        md += "---\n\n"
        md += "## Decision\n\n"
        md += "☐ **PROCEED** - Gate criteria met, continue to next phase\n\n"
        md += "☐ **CONDITIONAL** - Minor issues to address, proceed with monitoring\n\n"
        md += "☐ **HOLD** - Significant issues require resolution before proceeding\n\n"
        md += "☐ **CANCEL** - Project no longer viable\n\n"
        
        md += "**Decision Date:** ________________\n\n"
        md += "**Approved By:** ________________\n\n"
        md += "**Notes:**\n\n"
        md += "_" * 80 + "\n\n"
        md += "_" * 80 + "\n\n"
        
        return md
    
    def export_team_report(
        self,
        recommendation,
        output_path: str,
        format: str = 'pdf'
    ) -> str:
        """
        Export team composition report
        
        Args:
            recommendation: TeamRecommendation object
            output_path: Output file path
            format: 'pdf', 'word', or 'markdown'
        
        Returns:
            Output file path
        """
        # For now, export as markdown
        md = "# Team Composition Report\n\n"
        md += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n"
        md += "---\n\n"
        
        md += "## Recommendation Summary\n\n"
        md += f"- **Skill Match:** {recommendation.overall_skill_match:.1f}%\n"
        md += f"- **Team Size:** {recommendation.team_size_fte:.1f} FTE\n"
        md += f"- **Total Cost:** ${recommendation.total_cost:,.0f}\n"
        md += f"- **Predicted Performance:** {recommendation.predicted_performance:.1f}/100\n"
        md += f"- **Confidence:** {recommendation.confidence:.1f}%\n\n"
        
        md += "## Team Composition\n\n"
        for member in recommendation.team_members:
            md += f"### {member.person.name} ({member.person.role})\n"
            md += f"- **Allocation:** {member.allocation*100:.0f}%\n"
            md += f"- **Skill Match:** {member.skill_match_score:.0f}%\n"
            md += f"- **Rationale:** {member.rationale}\n\n"
        
        if recommendation.strengths:
            md += "## Strengths\n\n"
            for strength in recommendation.strengths:
                md += f"- ✅ {strength}\n"
            md += "\n"
        
        if recommendation.risk_factors:
            md += "## Risk Factors\n\n"
            for risk in recommendation.risk_factors:
                md += f"- ⚠️ {risk}\n"
            md += "\n"
        
        if recommendation.skill_gaps:
            md += "## Skill Gaps\n\n"
            for gap in recommendation.skill_gaps:
                md += f"- 🔴 {gap}\n"
            md += "\n"
        
        # Write to file
        with open(output_path, 'w') as f:
            f.write(md)
        
        return output_path


# Demo usage
if __name__ == '__main__':
    print("📄 Report Templates Module")
    print("=" * 80)
    print()
    print("This module provides enhanced reporting capabilities:")
    print("  ✅ PDF export (requires: pip install reportlab)")
    print("  ✅ Word export (requires: pip install python-docx)")
    print("  ✅ Governance gate reports")
    print("  ✅ Team composition reports")
    print()
    print("Usage:")
    print("  from report_templates import ReportExporter")
    print("  exporter = ReportExporter()")
    print("  exporter.export_to_pdf(plan, 'plan.pdf')")
    print("  exporter.export_to_word(plan, 'plan.docx')")
    print()
    print("Note: Install dependencies for full functionality:")
    print("  pip install reportlab python-docx")
    print()
    print("Markdown export works without additional dependencies.")
